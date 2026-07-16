"""Office Unstructured + text native parse (#280)."""

from __future__ import annotations

import io
import json

import pytest
from docx import Document

from path_graph.parsers.parse import (
    parse_non_pdf_to_blocks,
    parse_office_to_blocks,
    parse_text_to_blocks,
)
from path_graph.steps.ingest import ingest_raw_bytes
from path_graph.storage.blob import LocalBlobStore

from constants import PROJECT_ID


def _minimal_docx() -> bytes:
    doc = Document()
    doc.add_heading("Section Title", level=1)
    doc.add_paragraph("Office body paragraph for unstructured.")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@pytest.fixture
def local_store(tmp_path, monkeypatch):
    monkeypatch.setenv("PIPELINE_STORAGE_BACKEND", "local")
    monkeypatch.setenv("PIPELINE_STORAGE_DIR", str(tmp_path))
    return tmp_path


def test_parse_text_to_blocks():
    doc = parse_text_to_blocks(b"Hello\n\nWorld")
    assert doc["extractor"] == "text"
    assert [b["text"] for b in doc["blocks"]] == ["Hello", "World"]


def test_parse_office_docx_to_blocks():
    doc = parse_office_to_blocks(_minimal_docx(), "sample.docx")
    assert doc["extractor"] == "unstructured"
    assert doc["blocks"]
    texts = " ".join(b.get("text") or "" for b in doc["blocks"])
    assert "Section Title" in texts
    assert "Office body paragraph" in texts


def test_parse_non_pdf_routes_docx():
    doc = parse_non_pdf_to_blocks(_minimal_docx(), "a.docx")
    assert doc["extractor"] == "unstructured"


def test_ingest_docx_uses_unstructured_backend(local_store):
    meta = {
        "tenant": "dev",
        "document_id": "00000000-0000-0000-0000-000000000021",
        "content_hash": "docxhash",
        "source_id": "manual",
        "project_id": PROJECT_ID,
        "s3_raw_uri": "file://x",
        "filename": "sample.docx",
    }
    result = ingest_raw_bytes(_minimal_docx(), "sample.docx", "dev", "manual", meta)
    assert result["chunks"]
    assert result["parse_backend"] == "unstructured"
    store = LocalBlobStore(local_store)
    blocks = json.loads(
        store.get_bytes(f"parsed/dev/{meta['document_id']}/content.json")
    )
    assert blocks["extractor"] == "unstructured"
